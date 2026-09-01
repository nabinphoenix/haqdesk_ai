import logging
import re
import time
from typing import Optional, Literal
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.models import (
    Distance, VectorParams, PointStruct,
    Filter, FieldCondition, MatchValue,
    SearchRequest
)
from sqlalchemy.orm import Session
from app.core.config import settings
from app.models.knowledge import KnowledgeDocument, KnowledgeChunk
from app.models.message import Message
from app.models.business import Business

from app.services.nepali_normalizer import normalize_nepali_text, get_embedding_input
from app.services.qa_parser import parse_qa_pairs
from app.prompts.customer_reply_prompt import build_system_prompt
from app.services.llm_gateway import llm_gateway

logger = logging.getLogger("uvicorn")

# Embedding dimension — must match the model configured in settings
EMBEDDING_DIM = settings.EMBEDDING_DIM
CONFIDENCE_THRESHOLD = 0.45


_OWNER_CONTACT_TERMS = (
    "owner", "business owner", "malik", "malikko", "malik ko",
)
_CONTACT_TERMS = (
    "contact", "contact detail", "phone", "number", "mobile", "email",
    "sampark", "sampak", "bibaran", "detail",
)
_ROMANIZED_NEPALI_MARKERS = (
    "hajur", "tapai", "tapailai", "malai", "hamro", "cha", "chha", "xa",
    "xaina", "chaina", "garnu", "milxa", "bhayo", "ko", "yo", "thiyo",
)


def _is_private_owner_contact_request(text: str) -> bool:
    """Identify requests for an owner's personal contact details."""
    normalized = " ".join((text or "").lower().split())
    return bool(
        normalized
        and any(term in normalized for term in _OWNER_CONTACT_TERMS)
        and any(term in normalized for term in _CONTACT_TERMS)
    )


def _looks_romanized_nepali(text: str) -> bool:
    tokens = set(re.findall(r"[a-z]+", (text or "").lower()))
    return bool(tokens.intersection(_ROMANIZED_NEPALI_MARKERS))




def private_owner_contact_response(question: str) -> str | None:
    """Give a concise privacy-safe answer without inventing contact information."""
    if not _is_private_owner_contact_request(question):
        return None

    if _looks_romanized_nepali(question):
        return (
            "Hello,\n\nHami byabasaya malikko byaktigat samparka bibaran share garna "
            "sakdainau. Kripaya yahai message ma aafno prasna pathaunuhos; hamro support "
            "teamle sahayog garchha."
        )

    return (
        "Hello,\n\nWe cannot share the business owner's private contact details. "
        "Please send your question here and the support team will help."
    )



_LEXICAL_STOP_WORDS = frozenset({
    "a", "an", "the", "can", "could", "do", "does", "did", "is", "are", "am",
    "was", "were", "what", "how", "why", "where", "when", "who", "which", "would",
    "should", "may", "might", "i", "me", "my", "we", "us", "our", "you", "your",
    "they", "them", "this", "that", "these", "those", "to", "of", "for", "on", "in",
    "at", "from", "with", "through", "before", "after", "and", "or", "like", "please",
    "tell", "want", "need", "guys", "text", "customer", "customers",
})


def _lexical_tokens(value: str) -> set[str]:
    """Return small normalized keywords for a safe SQL retrieval fallback."""
    tokens = re.findall(r"[a-z0-9]+", (value or "").casefold())
    normalized = []
    for token in tokens:
        if token in _LEXICAL_STOP_WORDS:
            continue
        if token.endswith("ies") and len(token) > 4:
            token = f"{token[:-3]}y"
        elif token.endswith("ing") and len(token) > 5:
            token = token[:-3]
        elif token.endswith("ed") and len(token) > 4:
            token = token[:-2]
        elif token.endswith("s") and len(token) > 3:
            token = token[:-1]
        normalized.append(token)
    return set(normalized)


def _lexical_match_score(query: str, content: str) -> float:
    """Score query keyword coverage, prioritizing the stored FAQ question."""
    query_tokens = _lexical_tokens(query)
    if not query_tokens:
        return 0.0
    question_part = content.split("\nA:", 1)[0]
    question_tokens = _lexical_tokens(question_part)
    content_tokens = _lexical_tokens(content)
    question_overlap = len(query_tokens & question_tokens)
    content_overlap = len(query_tokens & content_tokens)
    query_size = len(query_tokens)
    weighted_score = (
        0.75 * question_overlap / query_size
        + 0.25 * content_overlap / query_size
    )
    if question_overlap and content_overlap == query_size:
        return max(weighted_score, 0.9)
    return weighted_score
class RAGService:
    """
    RAG Service with multilingual embedding (intfloat/multilingual-e5-small by default)
    & Qdrant Vector Store.
    """
    def __init__(self):
        self._embedder = None
        self._qdrant = None
        self._initialized_collections = set()

    @property
    def embedder(self):
        if self._embedder is None:
            model_name = settings.EMBEDDING_MODEL
            logger.info("[RAG] Loading %s embedding model...", model_name)
            self._embedder = SentenceTransformer(model_name)
            logger.info("[RAG] %s embedding model loaded successfully.", model_name)
        return self._embedder

    @property
    def qdrant(self):
        if self._qdrant is None:
            try:
                client = QdrantClient(
                    host=settings.QDRANT_HOST,
                    port=settings.QDRANT_PORT,
                    timeout=3.0
                )
                # Test connection
                client.get_collections()
                self._qdrant = client
                logger.info(f"[RAG] Connected to Qdrant server at {settings.QDRANT_HOST}:{settings.QDRANT_PORT}")
            except Exception as e:
                if not settings.QDRANT_ALLOW_LOCAL_FALLBACK:
                    raise RuntimeError('Shared Qdrant is unavailable and local fallback is disabled.') from e
                logger.warning('[RAG] Remote Qdrant unavailable; using the development-only local index: %s', e)
                self._qdrant = QdrantClient(path=settings.VECTOR_DB_PATH)

        return self._qdrant

    @staticmethod
    def collection_name_for_business(business_id: int) -> str:
        """Derive the only allowed Qdrant collection name for a tenant."""
        if not isinstance(business_id, int) or isinstance(business_id, bool) or business_id <= 0:
            raise ValueError("A positive integer business_id is required")
        prefix = settings.QDRANT_COLLECTION_PREFIX.strip().strip("_")
        if not prefix:
            raise ValueError("QDRANT_COLLECTION_PREFIX cannot be blank")
        return f"{prefix}_{business_id}"

    def _collection_exists(self, business_id: int) -> bool:
        collection_name = self.collection_name_for_business(business_id)
        if collection_name in self._initialized_collections:
            return True
        exists = self.qdrant.collection_exists(collection_name=collection_name)
        if exists:
            self._initialized_collections.add(collection_name)
        return exists

    def _ensure_business_collection(self, business_id: int) -> str:
        """Lazily create one non-destructive collection for a business."""
        collection_name = self.collection_name_for_business(business_id)
        if not self._collection_exists(business_id):
            self.qdrant.create_collection(
                collection_name=collection_name,
                vectors_config=VectorParams(
                    size=EMBEDDING_DIM,
                    distance=Distance.COSINE,
                ),
            )
            logger.info("[RAG] Created business collection '%s'", collection_name)

        info = self.qdrant.get_collection(collection_name=collection_name)
        vectors = info.config.params.vectors
        current_size = vectors.size if hasattr(vectors, "size") else None
        if current_size != EMBEDDING_DIM:
            raise RuntimeError(
                f"Qdrant collection '{collection_name}' has vector size "
                f"{current_size}; expected {EMBEDDING_DIM}. Refusing destructive recreation."
            )
        self._initialized_collections.add(collection_name)
        return collection_name

    def embed_text(self, text: str, prefix: str = "query: ") -> list[float]:
        """Embed a single text after Romanized Nepali normalization.

        Args:
            text: Raw input text.
            prefix: E5-family models require a task prefix.
                    Use "query: " for search queries (default),
                    "passage: " for document chunks during ingestion.
        """
        normalized_input = get_embedding_input(text)
        return self.embedder.encode(f"{prefix}{normalized_input}").tolist()

    def embed_batch(self, texts: list[str], prefix: str = "passage: ") -> list[list[float]]:
        """Embed multiple texts after Romanized Nepali normalization.

        Args:
            texts: List of raw input texts.
            prefix: E5-family models require a task prefix.
                    Use "passage: " for document chunks (default),
                    "query: " for batch search queries.
        """
        normalized_inputs = [f"{prefix}{get_embedding_input(t)}" for t in texts]
        return self.embedder.encode(normalized_inputs).tolist()

    @staticmethod
    def extract_document_pages(file_path: str) -> tuple[list[tuple[int, str]], str]:
        '''Extract text with a parser appropriate for the uploaded file type.'''
        from pathlib import Path

        suffix = Path(file_path).suffix.lower()
        pages: list[tuple[int, str]] = []
        if suffix == '.pdf':
            import fitz
            with fitz.open(file_path) as document:
                for page_number, page in enumerate(document, start=1):
                    pages.append((page_number, page.get_text() or ''))
        elif suffix == '.docx':
            from docx import Document
            document = Document(file_path)
            lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        lines.append(' | '.join(cells))
            pages = [(1, '\n'.join(lines))]
        elif suffix == '.txt':
            raw = Path(file_path).read_text(encoding='utf-8', errors='replace')
            pages = [(1, raw)]
        else:
            display_suffix = suffix or 'unknown'
            raise ValueError(f'Unsupported document type: {display_suffix}')

        full_text = '\n'.join(text for _, text in pages).strip()
        if len(full_text) > settings.KNOWLEDGE_MAX_EXTRACTED_CHARACTERS:
            raise ValueError('Extracted document text exceeds the configured limit.')
        return pages, full_text
    def ingest_document(
        self,
        file_path: str,
        filename: str,
        document_id: int,
        business_id: int,
        db: Session,
        source_type: str = "upload"
    ):
        """Parse document into Q&A pairs (or fallback chunks), embed chunks, store in Qdrant and PostgreSQL."""
        logger.info(f"[RAG] Ingesting document: {filename} (ID={document_id})")

        try:
            full_text_by_page, full_raw_text = self.extract_document_pages(file_path)

            # Attempt Fix 2: Q&A-pair parsing
            parsed_qa = parse_qa_pairs(full_raw_text)
            chunks = []

            if parsed_qa:
                logger.info(f"[RAG] Document identified as Q&A dataset with {len(parsed_qa)} pairs.")
                # Map extracted Q&A pairs back to page numbers where possible
                for qa in parsed_qa:
                    matched_page = 1
                    for pg_num, pg_txt in full_text_by_page:
                        if qa["question"][:30] in pg_txt:
                            matched_page = pg_num
                            break
                    chunks.append({
                        "content": qa["content"],
                        "page_number": matched_page,
                    })
            else:
                logger.info(f"[RAG] No Q&A structure found. Falling back to sliding-window chunking.")
                for pg_num, pg_txt in full_text_by_page:
                    words = pg_txt.split()
                    chunk_size = 500
                    overlap = 50
                    i = 0
                    while i < len(words):
                        chunk_words = words[i:i + chunk_size]
                        chunk_text = " ".join(chunk_words).strip()
                        if chunk_text:
                            chunks.append({
                                "content": chunk_text,
                                "page_number": pg_num,
                            })
                        i += chunk_size - overlap

            logger.info(f"[RAG] Split document into {len(chunks)} total chunks.")

            if not chunks:
                raise ValueError("No text extracted from document")

            # Embed all chunks
            embeddings = self.embed_batch([c["content"] for c in chunks])
            qdrant_points = []

            for idx, (chunk_data, vector) in enumerate(zip(chunks, embeddings)):
                # Fix 3: Store metadata in PostgreSQL without embedding column
                db_chunk = KnowledgeChunk(
                    business_id=business_id,
                    document_id=document_id,
                    content=chunk_data["content"],
                    page_number=chunk_data["page_number"],
                )
                db.add(db_chunk)
                db.flush()  # Gets db_chunk.id

                # Qdrant is the sole vector store
                qdrant_points.append(
                    PointStruct(
                        id=db_chunk.id,
                        vector=vector,
                        payload={
                            "business_id": business_id,
                            "document_id": document_id,
                            "chunk_id": db_chunk.id,
                            "content": chunk_data["content"],
                            "page_number": chunk_data["page_number"],
                            "filename": filename,
                            "source_type": source_type,

                        }
                    )
                )

            # Keep SQL rows and vectors in sync. Qdrant has no SQL transaction,
            # so remove newly written points if either side fails.
            collection_name = self._ensure_business_collection(business_id)
            point_ids = [point.id for point in qdrant_points]
            try:
                self.qdrant.upsert(
                    collection_name=collection_name,
                    points=qdrant_points,
                )
                db.commit()
            except Exception:
                db.rollback()
                if point_ids:
                    try:
                        from qdrant_client.models import PointIdsList
                        self.qdrant.delete(
                            collection_name=collection_name,
                            points_selector=PointIdsList(points=point_ids),
                        )
                    except Exception:
                        logger.exception('[RAG] Failed to clean up partial Qdrant points')
                raise

            logger.info(f"[RAG] Successfully stored {len(chunks)} Q&A chunks for doc {document_id} in Qdrant & DB")
            return len(chunks)

        except Exception as e:
            db.rollback()
            logger.error(f"[RAG] Document ingestion failed: {e}")
            raise e

    def _retrieve_lexical_chunks(
        self,
        query_text: str,
        business_id: int,
        db: Optional[Session],
        top_k: int,
    ) -> list[dict]:
        """Retrieve tenant-owned FAQ rows when vector search is unavailable or weak."""
        if db is None:
            return []

        rows = (
            db.query(KnowledgeChunk, KnowledgeDocument)
            .join(KnowledgeDocument, KnowledgeDocument.id == KnowledgeChunk.document_id)
            .filter(
                KnowledgeChunk.business_id == business_id,
                KnowledgeDocument.business_id == business_id,
                KnowledgeDocument.status == "ready",
            )
            .all()
        )
        candidates = []
        for chunk, document in rows:
            score = _lexical_match_score(query_text, chunk.content)
            if score <= 0:
                continue
            candidates.append({
                "content": chunk.content,
                "similarity": score,
                "lexical_similarity": score,
                "page_number": chunk.page_number,
                "filename": document.filename or "",
                "chunk_id": chunk.id,
                "document_id": chunk.document_id,
                "source_type": document.source_type or "upload",
            })

        candidates.sort(
            key=lambda item: (item["similarity"], len(item["content"])),
            reverse=True,
        )
        if candidates:
            unique_candidates = []
            seen_questions = set()
            for item in candidates:
                question_key = item["content"].split("\nA:", 1)[0].strip().casefold()
                if question_key in seen_questions:
                    continue
                seen_questions.add(question_key)
                unique_candidates.append(item)
            candidates = unique_candidates
            best_score = candidates[0]["similarity"]
            minimum_score = max(CONFIDENCE_THRESHOLD, best_score - 0.2)
            candidates = [
                item for item in candidates if item["similarity"] >= minimum_score
            ]
        return candidates[:top_k]

    def retrieve_chunks(
        self,
        query_text: str,
        business_id: int,
        top_k: int = 5,
        db: Optional[Session] = None,
    ) -> list[dict]:
        """Search vectors and fall back to tenant-scoped lexical FAQ retrieval."""
        chunks = []
        try:
            if not self._collection_exists(business_id):
                logger.info("[RAG] No knowledge collection exists for business %s", business_id)
            else:
                collection_name = self.collection_name_for_business(business_id)
                query_embedding = self.embed_text(query_text)
                query_filter = Filter(
                    must=[
                        FieldCondition(
                            key="business_id",
                            match=MatchValue(value=business_id),
                        )
                    ]
                )
                candidate_limit = max(top_k, top_k * 4)

                if hasattr(self.qdrant, "search"):
                    results = self.qdrant.search(
                        collection_name=collection_name,
                        query_vector=query_embedding,
                        query_filter=query_filter,
                        limit=candidate_limit,
                        with_payload=True,
                    )
                else:
                    response = self.qdrant.query_points(
                        collection_name=collection_name,
                        query=query_embedding,
                        query_filter=query_filter,
                        limit=candidate_limit,
                        with_payload=True,
                    )
                    results = response.points

                for result in results:
                    payload = result.payload or {}
                    chunks.append({
                        "content": payload.get("content", ""),
                        "similarity": result.score,
                        "page_number": payload.get("page_number"),
                        "filename": payload.get("filename", ""),
                        "chunk_id": payload.get("chunk_id"),
                        "document_id": payload.get("document_id"),
                        "source_type": payload.get("source_type", "upload"),
                    })
        except Exception as e:
            logger.warning(
                "[RAG] Qdrant search unavailable (%s). Trying tenant-scoped FAQ retrieval.",
                e,
            )

        vector_score = max(
            (float(chunk.get("similarity") or 0) for chunk in chunks),
            default=0.0,
        )
        if db is not None and (not chunks or vector_score < CONFIDENCE_THRESHOLD):
            try:
                lexical_chunks = self._retrieve_lexical_chunks(
                    query_text, business_id, db, top_k
                )
                if lexical_chunks:
                    by_chunk_id = {
                        chunk.get("chunk_id"): chunk
                        for chunk in chunks
                        if chunk.get("chunk_id") is not None
                    }
                    for lexical in lexical_chunks:
                        existing = by_chunk_id.get(lexical["chunk_id"])
                        if existing is None:
                            chunks.append(lexical)
                            continue
                        existing["lexical_similarity"] = lexical["lexical_similarity"]
                        existing["similarity"] = max(
                            float(existing.get("similarity") or 0),
                            lexical["lexical_similarity"],
                        )
                    chunks.sort(
                        key=lambda item: (
                            float(item.get("similarity") or 0),
                            len(item.get("content") or ""),
                        ),
                        reverse=True,
                    )
                    return chunks[:top_k]
            except Exception as e:
                logger.warning("[RAG] Tenant-scoped FAQ retrieval failed: %s", e)

        return chunks[:top_k]

    async def query(
        self,
        question: str,
        business_id: int,
        db: Session,
        conversation_id: Optional[int] = None,
        current_message_id: Optional[int] = None,
        mode: Literal["auto", "review"] = "review",
        top_k: int = 5,
        confidence_threshold: float = CONFIDENCE_THRESHOLD,
        language: Optional[str] = None,
        sentiment: Optional[str] = None,
        platform: Optional[str] = None,
        customer_name: Optional[str] = None,
        business_name: Optional[str] = None,
        **kwargs
    ) -> Optional[dict]:
        """Full RAG pipeline: retrieve → generate → return answer with context memory."""
        start_time = time.time()
        resolved_business_name = (business_name or "").strip()

        try:
            direct_response = private_owner_contact_response(question)
            if direct_response:
                detected_lang = (
                    "romanized_nepali"
                    if _looks_romanized_nepali(question)
                    else (language or "english")
                )
                return {
                    "answer": direct_response,
                    "confidence": 0.0,
                    "sources": [],
                    "chunks_used": 0,
                    "source_details": [],
                    "grounded": False,
                    "language_detected": detected_lang,
                    "metadata": {
                        "model": "policy-rule",
                        "provider": "internal",
                        "fallback_used": False,
                        "attempts": 0,
                        "latency_ms": round((time.time() - start_time) * 1000, 2),
                        "grounded": False,
                        "direct_response": "private_owner_contact",
                    },
                }
            from app.services.context_translator import context_translator

            # 1. Dynamically analyze language and get English translation for search
            analysis = await context_translator.analyze_and_translate(question)
            english_question = analysis["english_translation"]
            detected_lang = language or analysis["detected_language"]

            if not resolved_business_name and db:
                business = db.query(Business).filter(
                    Business.id == business_id
                ).first()
                resolved_business_name = (business.name or "").strip() if business else ""
            resolved_business_name = resolved_business_name or "Support Team"

            # 2. Retrieve relevant chunks using the English question
            chunks = self.retrieve_chunks(english_question, business_id, top_k, db=db)
            top_score = chunks[0]["similarity"] if chunks else 0.0

            sources = []
            source_details = []
            grounded = bool(chunks and top_score >= confidence_threshold)
            if grounded:
                logger.info(f'[RAG] Retrieved {len(chunks)} chunks. Top score: {top_score:.3f}')
                context = '\n\n---\n\n'.join([
                    f'[Page {c["page_number"]}] {c["content"]}'
                    for c in chunks
                ])
                source_details = [
                    {
                        'filename': c.get('filename', ''),
                        'page_number': c.get('page_number'),
                        'chunk_id': c.get('chunk_id'),
                        'similarity': round(float(c.get('similarity') or 0), 4),
                        'source_type': c.get('source_type', 'upload'),
                    }
                    for c in chunks
                ]
                sources = list(dict.fromkeys(
                    detail['filename'] for detail in source_details if detail['filename']
                ))
            else:
                logger.info(
                    f'[RAG] Low confidence ({top_score:.3f}) or no chunks found. '
                    'Refusing cross-business/general policy assumptions.'
                )
                context = (
                    'No sufficiently relevant business document was found. '
                    'Do not invent an answer, use a different business policy, or promise a future reply. '
                    'Briefly say the information cannot be confirmed here and ask the customer to '
                    'send the relevant details in this chat.'
                )

            # 2. Build system prompt using unified prompt builder (Fix 4)
            system_prompt = build_system_prompt(
                context=context,
                mode=mode,
                language=detected_lang,
                sentiment=sentiment,
                platform=platform,
                customer_name=customer_name,
                business_name=resolved_business_name,
            )

            # 3. Retrieve conversation history with current_message_id exclusion (Fix 1)
            past_messages = []
            if db and conversation_id:
                history_query = db.query(Message).filter(
                    Message.conversation_id == conversation_id
                )
                if current_message_id is not None:
                    history_query = history_query.filter(Message.id != current_message_id)

                db_messages = history_query.order_by(Message.timestamp.desc()).limit(10).all()
                db_messages.reverse()

                for m in db_messages:
                    if not m.content:
                        continue
                    role = "user" if m.sender_type == "customer" else "assistant"
                    past_messages.append({"role": role, "content": m.content})

            messages = [{"role": "system", "content": system_prompt}]

            for pm in past_messages:
                messages.append(pm)

            # Unconditionally append current question as final user message (Fix 1)
            messages.append({"role": "user", "content": question})

            result = await llm_gateway.complete(messages=messages, max_tokens=2000)

            if not result or not result.get("content"):
                logger.warning("[RAG] LLM returned empty response")
                raise RuntimeError("LLM returned an empty response")

            generated_answer = result["content"]

            # 4. Translate the generated response back to the customer's language style
            final_answer = await context_translator.translate_to_target_language(generated_answer, detected_lang)

            latency_ms = round((time.time() - start_time) * 1000, 2)

            return {
                "answer": final_answer,
                "confidence": top_score if grounded else 0.0,
                "sources": sources,
                "chunks_used": len(chunks) if grounded else 0,
                "source_details": source_details,
                "grounded": grounded,
                "language_detected": detected_lang,
                "metadata": {
                    "model": result.get("model", "unknown"),
                    "provider": result.get("provider", "unknown"),
                    "fallback_used": result.get("fallback_used", False),
                    "attempts": result.get("attempts", 1),
                    "latency_ms": latency_ms,
                    "grounded": grounded,
                    "confidence_threshold": confidence_threshold,
                    "source_details": source_details,
                }
            }

        except Exception as e:
            import traceback
            logger.error(f"[RAG] Query failed:\n{traceback.format_exc()}")
            lang = language or "english"
            if not resolved_business_name and db:
                try:
                    business = db.query(Business).filter(
                        Business.id == business_id
                    ).first()
                    resolved_business_name = (
                        (business.name or "").strip() if business else ""
                    )
                except Exception:
                    resolved_business_name = ""
            resolved_business_name = resolved_business_name or "Support Team"
            fallback_answer = (
                f"Namaste! {resolved_business_name} support ma swagat cha. "
                "Ahile tapai ko request process garna samasya bhairako cha. "
                "Hamro team ko sadasya le chhittai follow up garnuhunecha."
                if lang == "romanized_nepali"
                else f"Hello! Welcome to {resolved_business_name} support. "
                "We're having trouble processing your request. "
                "A team member will follow up shortly."
            )
            return {
                "answer": fallback_answer,
                "confidence": 0.0,
                "sources": [],
                "chunks_used": 0,
                "source_details": [],
                "grounded": False,
                "language_detected": lang,
                "metadata": {"fallback_used": True, "grounded": False, "error": str(e)}
            }

    def delete_business_collection(self, business_id: int) -> bool:
        '''Permanently remove the isolated vector collection for one business.'''
        if not self._collection_exists(business_id):
            return False
        collection_name = self.collection_name_for_business(business_id)
        self.qdrant.delete_collection(collection_name=collection_name)
        self._initialized_collections.discard(collection_name)
        logger.info('[RAG] Deleted business collection %s', collection_name)
        return True

    def delete_document_chunks(self, document_id: int, business_id: int, db: Session):
        """Delete all chunks for a document from both Qdrant and PostgreSQL."""
        chunks = db.query(KnowledgeChunk).filter(
            KnowledgeChunk.document_id == document_id,
            KnowledgeChunk.business_id == business_id,
        ).all()

        chunk_ids = [c.id for c in chunks]

        if chunk_ids and self._collection_exists(business_id):
            try:
                from qdrant_client.models import PointIdsList
                self.qdrant.delete(
                    collection_name=self.collection_name_for_business(business_id),
                    points_selector=PointIdsList(points=chunk_ids)
                )
                logger.info(f"[RAG] Deleted {len(chunk_ids)} points from Qdrant")
            except Exception as e:
                logger.error(f"[RAG] Failed to delete from Qdrant: {e}")
                raise

        db.query(KnowledgeChunk).filter(
            KnowledgeChunk.document_id == document_id,
            KnowledgeChunk.business_id == business_id,
        ).delete()
        db.commit()

    def update_chunk_embedding(
        self, chunk_id: int, new_content: str, business_id: int, db: Session
    ):
        """Re-embed and update a single chunk in Qdrant and update content in PostgreSQL."""
        chunk = db.query(KnowledgeChunk).filter(
            KnowledgeChunk.id == chunk_id,
            KnowledgeChunk.business_id == business_id,
        ).first()
        if not chunk:
            raise ValueError(f"Chunk {chunk_id} not found")

        new_embedding = self.embed_text(new_content, prefix="passage: ")

        # Update PostgreSQL content after vector indexing succeeds
        chunk.content = new_content

        document = db.query(KnowledgeDocument).filter(
            KnowledgeDocument.id == chunk.document_id,
            KnowledgeDocument.business_id == business_id,
        ).first()

        # Update Qdrant (sole vector store)
        collection_name = self._ensure_business_collection(business_id)
        self.qdrant.upsert(
            collection_name=collection_name,
            points=[
                PointStruct(
                    id=chunk_id,
                    vector=new_embedding,
                    payload={
                        "business_id": chunk.business_id,
                        "document_id": chunk.document_id,
                        "chunk_id": chunk_id,
                        "content": new_content,
                        "page_number": chunk.page_number,
                        "filename": document.filename if document else "",
                        "source_type": document.source_type if document else "upload",
                    }
                )
            ]
        )
        if document and document.status == "draft":
            document.status = "ready"
        db.commit()
        logger.info(f"[RAG] Chunk {chunk_id} updated in Qdrant and PostgreSQL")


rag_service = RAGService()
